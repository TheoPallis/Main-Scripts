import glob
import os
import shutil
import subprocess
import docx
from docx import Document
import pandas as pd
import numpy as np
import openpyxl
import re
import xlsxwriter
import xlrd
import logging
from openpyxl import load_workbook
import streamlit as st  
from streamlit_extras.no_default_selectbox import selectbox as sb
from streamlit_extras.stateful_button import button

import tkinter as tk
from tkinter import filedialog as fd
# from olografos_clean import *
from openpyxl.styles import Font, Color, PatternFill
import logging
logging.basicConfig(filename='app.log', filemode='w', format='%(name)s - %(levelname)s - %(message)s')
# import pydocx
output = r'C:\Users\pallist\Desktop\ΤΡΕΧΟΝΤΑ\Testing Folder\TFT output'
os.chdir(output)                 


def format_df(file) :
    workbook = openpyxl.load_workbook(file)
    for worksheet in workbook:
        font = Font(color='FFFFFF', bold=True)
        fill = PatternFill(start_color='5552A2', end_color='5552A2', fill_type='solid')
        for cell in worksheet[1]:
            cell.font = font
            cell.fill = fill
        for column in worksheet.columns:
            worksheet.column_dimensions[openpyxl.utils.get_column_letter(column[0].column)].width = 50

        
    workbook.save(file)


st.sidebar.header("Options")
sidebar_options = [
'Αρχική Σελίδα', 
'1) Συνένωση Excel',
'2) Συμπλήρωση Excel βάσει δύο στηλών',
'3) Επεξεργασία & Συνένωση Δεδομένων',
'4) Διαχωρισμός',
'5) Ένορκες Βεβαιώσεις',
'6) Μαζικές εργασίες εγγράφων'
]
sidebar = st.sidebar.selectbox( "Επιλογή εργασιών :", sidebar_options)





if sidebar =='3) Επεξεργασία & Συνένωση Δεδομένων':

# 1a. Αν επιλεγεί στήλη αποθηκεύεται στο selected_columns_dict με τη μορφή : 
#     selected_columns['user_defined_name'] = selected_column
# 1b. Επιλογή 4 τιμών στήλης και εισαγωγή σε λίστα 1
# 1c. Επιλογή 4 λεκτικών και εισαγωγή σε λίστα 2
# 1d. Χρήση replace τιμών λίστας 1 με τιμές λίστας 2
# 1e. Αφαίρεση κενών τιμών στο selected_columns_dict
# 2a. Δημιουργία dict για χειρισμό εξαιρέσεων πόλεων  
# 2b. Δημιουργία 2ου dict για χειρισμό εξαιρέσεων πόλεων (άγιος/νέος)  
# 2c. Δημιουργία λίστας με όλες τις καταλήξεις εταιριών για αναγνώριση περίπτωησς εταρείας 
# 3a. Του       : Αν το όνομα τελειώνει σε Α/Η -> Της, ΟΣ/ΑΣ/ΉΣ/υπόλοιπα -> Του 
# 3b. Τον       : Αν το όνομα τελειώνει σε Α/Η -> Την, ΟΣ/ΑΣ/ΗΣ/υπόλοιπα -> Τον
# 3d. Αιτιατική : Αν το όνομα τελειώνει σε Α/Η/υπόλοιπα -> όνομα, ΟΣ/ΑΣ/ΗΣ -> όνομα χωρίς το ς
# 3e. Γενική    : Αν το όνομα τελειώνει σε Α/Η/υπόλοιπα -> όνομα + ς, ΑΣ/ΗΣ -> όνομα χωρίς το ς, ΟΣ -> ΟΥ
# 3f. Πατρώνυμο : Αν το όνομα τελειώνει σε ΑΣ/ΗΣ -> όνομα χωρίς το ς, ΟΣ -> ΟΥ
# 4a. Πόλη  (αν επιλεγεί) : Αν η πόλη βρίσκεται στο 2a ή στο 2b,αντικατάσταση πόλης-key με πόλη=value, ΑΣ->Α, ΕΣ/ΟΙ -> ΩΝ, Α/Η-> ΑΣ/ΗΣ, Ο,Ι->ΟΥ                                                
#  Να μετατρέψω σε κεφαλαία❗


    # Global variable to store selected columns
 

    uploaded_file = st.file_uploader("Επιλογή αρχείου Excel")
    if uploaded_file is not None:
        arxiko_df=pd.read_excel(uploaded_file,sheet_name= None,dtype=str)
        selected_sheet = st.selectbox("Επιλογή Φύλλου :",options = arxiko_df.keys())
        df = arxiko_df[selected_sheet].fillna("")
        df_obj = df.select_dtypes(['object'])
        df[df_obj.columns] = df_obj.apply(lambda x: x.str.strip())
        selected_columns ={}

        if st.button('Προβολή φύλλου') :
            df
                
        options = st.selectbox("Επιλογές", ['Αυτόματη Εύρεση Στηλών Non Frontier','Αυτόματη Εύρεση Στηλών Frontier', 'Χειροκίνητη Εύρεση Στηλών'])

        custom_city_gen = 'κατοίκου '
        custom_street_gen =  'οδός '
        custom_TK_gen =  'ΤΚ '
        custom_ΑΦΜ_gen = ', με ΑΦΜ '
        custom_city_ait =  'κάτοικο '
        custom_street_ait = custom_street_gen
        custom_TK_ait =   custom_TK_gen
        custom_ΑΦΜ_ait =    custom_ΑΦΜ_gen

        if options == 'Χειροκίνητη Εύρεση Στηλών':
            col1, col2,col3= st.columns(3)
            with col1:
                st.subheader("Κύρια Στοιχεία")
                Όνομα = sb("Name", options=df.columns)
                if Όνομα :
                    selected_columns['Όνομα'] = Όνομα            
                # surname
                Επώνυμο = sb("Surname", options=df.columns)
                if Επώνυμο :
                    selected_columns['Επώνυμο'] = Επώνυμο
                # fathername
                Πατρώνυμο = sb("Fathername", options=df.columns)
                if Πατρώνυμο :
                    selected_columns['Πατρώνυμο'] = Πατρώνυμο
                    df['Πατρώνυμο'] = df[Πατρώνυμο] 
                # city
                Πόλη = sb("City", options=df.columns)
                if Πόλη :
                    selected_columns['Πόλη'] = Πόλη
                    df['City'] = df[Πόλη]
                # street
                Οδός = sb("Street", options=df.columns)
                if Οδός :
                    selected_columns['Οδός'] = Οδός
                    df['Οδός'] = df[Οδός]
                # number
                ΤΚ = sb("ΤΚ", options=df.columns)
                if ΤΚ :
                    selected_columns['TK'] = ΤΚ
                    df['ΤΚ'] = df[ΤΚ]
                # afm
                ΑΦΜ = sb("ΑΦΜ", options=df.columns)
                if ΑΦΜ :
                    selected_columns['ΑΦΜ'] = ΑΦΜ
                    df['ΑΦΜ'] = df[ΑΦΜ]
                # sxesi
                Σχέση = sb("Σχέση", options=df.columns)
                if Σχέση :
                    selected_columns['Σχέση'] = Σχέση
                    df['Σχέση'] = df[Σχέση]


                if st.button('Ορισμός λεκτικών' ):
                    with col3 :
                        st.subheader('Ορισμός λεκτικών')
                        custom_city_gen = st.text_input('Λεκτικό για κάτοικο (γενική)', ', κατοίκου ')
                        custom_street_gen = st.text_input ('Λεκτικό για οδό (γενική)', 'οδός')
                        custom_TK_gen = st.text_input ('Λεκτικό για ΤΚ (γενική) ', 'ΤΚ ')
                        custom_ΑΦΜ_gen = st.text_input ('Λεκτικό για ΑΦΜ (γενική)', ',με ΑΦΜ ')

                        custom_city_ait = st.text_input('Λεκτικό για κάτοικο (αιτιατική)', ', κάτοικο ')
                        custom_street_ait = st.text_input ('Λεκτικό για οδό (αιτιατική)', 'οδός')
                        custom_TK_ait = st.text_input ('Λεκτικό για ΤΚ (αιτιατική)', 'ΤΚ ')
                        custom_ΑΦΜ_ait = st.text_input ('Λεκτικό για ΑΦΜ (αιτιατική)', ',με ΑΦΜ ')
                        st.button('Εισαγωγή λεκτικών') # Πρέπει να πατηθεί κουμπί για να κρύψουμε στήλη
            with col2:
                st.subheader("Έξτρα Στοιχεία")
                # contract= sb("Contract (for pools)", options=df.columns)
                
                #amount
                Ποσό = sb("Ποσό", options=df.columns)
                if Ποσό :
                    selected_columns['Ποσό'] = Ποσό
                    df['Ποσό'] = df[Ποσό]
                
                # amount2
                Αριθμός = sb("Αριθμός", options=df.columns)
                

            if st.button("Προβολή επιλεχθέντων στηλών"):
                with col3 :
                    selected_columns = {k: v for k, v in selected_columns.items() if v is not None}
                    selected_columns 
                    # Create a new dataframe with only the selected columns
                # selected_df = df[selected_columns]
            
        elif options == 'Αυτόματη Εύρεση Στηλών Non Frontier' :

                
            selected_columns = {
                                "Όνομα":"total_main CUST FIRST NAME",
                                "Επώνυμο":"total_main CUST LAST NAME",
                                "Πατρώνυμο":"total_main CUST FATH NAME",
                                "Πόλη":"total_CITY",
                                "Οδός":"total_STREET",
                                "TK":"total_POSTCODE",
                                "ΑΦΜ":"total_main CUST AFM",
                                "Ποσό":"Denounced Amount",                                  
                                "Contract": "CASE CONTR NUM",
                                "Daneio" :"CASE_BU",
                                "SPV" : "ΧΑΡΤΟΦΥΛΑΚΙΟ",
                                "Σχέση" : "RT DESC"
        }
        
            Όνομα="total_main CUST FIRST NAME"
            Επώνυμο="total_main CUST LAST NAME"
            Πατρώνυμο="total_main CUST FATH NAME"
            Πόλη="total_CITY"
            Οδός="total_STREET"
            ΤΚ="total_POSTCODE"
            ΑΦΜ="total_main CUST AFM"
            Ποσό="Denounced Amount"                                  
            contract= "CASE CONTR NUM"
            daneio ="CASE_BU"
            SPV = "ΧΑΡΤΟΦΥΛΑΚΙΟ"
            Σχέση = "RT DESC"
            df['Οδός'] = df[Οδός]
            df['ΤΚ'] = df[ΤΚ]
            df['Πατρώνυμο'] = df[Πατρώνυμο]
            df['Πόλη'] =  df[Πόλη]
            df['Σχέση'] = df['RT DESC']
            df['ΑΦΜ'] = df[ΑΦΜ]

            

        elif options == 'Αυτόματη Εύρεση Στηλών Frontier' :
                selected_columns = {
                                    "Όνομα":"ΟΝΟΜΑ",
                                    "Επώνυμο":"ΕΠΩΝΥΜΟ",
                                    "Πατρώνυμο":"Father Name",
                                    "Πόλη":"Πόλη",
                                    "Οδός":"Οδός",
                                    "TK":"Τ.Κ.",
                                    "ΑΦΜ":"VAT Number",
                                    "Ποσό":"Denounced Amount",                                  
                                    "Contract": "Υπόθεση",
                                    "daneio" :"ΠΡΟΙΟΝ",
                                    "Σχέση" : "Σχέση",
                                    "Ποσό" :"Denounced Amount"  
                              
                }
            
                Όνομα = "ΟΝΟΜΑ"
                Επώνυμο = "ΕΠΩΝΥΜΟ"
                Πατρώνυμο = "ΠΑΤΡΩΝΥΜΟ"
                Πόλη = "Πόλη"
                Οδός = "Οδός"
                TK = "Τ.Κ."
                ΑΦΜ = "VAT Number"
                Ποσό = "Denounced Amount"                                  
                Contract =  "Υπόθεση"
                daneio ="ΠΡΟΙΟΝ"
                Σχέση = "Σχέση"
                Ποσό="Denounced Amount"  
                df['Σχέση'] = df[Σχέση]
    
        female_exceptions = ['ΕΛΙΣΑΒΕΤ']

        keys = ["ΒΡΙΛΗΣΙΑ","ΒΡΙΛΗΣΣΙΑ","ΙΛΙΟΝ","ΠΑΝΟΡΑΜΑ","ΣΕΡΡΕΣ","ΚΙΛΚΙΣ","ΑΜΠΕΛΑΚΙΑ","ΑΡΓΟΣ","ΧΑΝΙΑ","ΦΑΡΣΑΛΑ","ΙΩΑΝΝΙΝΑ","ΣΕΠΟΛΙΑ","ΤΡΙΚΑΛΑ","ΑΝΩ ΛΙΟΣΙΑ","ΚΑΛΥΒΙΑ","ΜΕΓΑΡΑ","ΝΈΟ ΗΡΑΚΛΕΙΟ","ΠΑΛΑΙΟ ΦΑΛΗΡΟ","ΝΕΑ ΙΩΝΙΑ","ΑΓΙΟΙ ΑΝΑΡΓΥΡΟΙ","ΑΓΙΑ ΠΑΡΑΣΚΕΥΗ"]

        values = ["ΒΡΙΛΗΣΙΩΝ","ΒΡΙΛΗΣΣΙΩΝ","ΙΛΙΟΥ","ΧΑΝΙΩΝ","ΠΑΝΟΡΑΜΑΤΟΣ","ΦΑΡΣΑΛΩΝ","ΑΜΠΕΛΑΚΙΩΝ","ΑΡΓΟΥΣ","ΙΩΑΝΝΙΝΩΝ","ΣΕΠΟΛΙΩΝ","ΤΡΙΚΑΛΩΝ","ΑΝΩ ΛΙΟΣΙΩΝ","ΚΑΛΥΒΙΩΝ","ΜΕΓΑΡΩΝ","ΝΕΟΥ ΗΡΑΚΛΕΙΟΥ","ΠΑΛΑΙΟΥ ΦΑΛΗΡΟΥ","ΝΕΑΣ ΙΩΝΙΑΣ","ΑΓΙΩΝ ΑΝΑΡΓΥΡΩΝ","ΑΓΙΑΣ ΠΑΡΑΣΚΕΥΗΣ",]

        exception_dict = dict(zip(keys, values))

        replacements = {
            "ΝΕΟΣ": "ΝΕΟΥ ",
            "ΑΓΙΟΣ ": "ΑΓΙΟΥ ",
            "ΝΕΑ": "ΝΕΑΣ ",
            "ΑΓΙΑ": "ΑΓΙΑΣ ",
            "ΑΓΙΟ": "ΑΓΙΟΥ ",
            "ΝΕΟ": "ΝΕΟΥ ",
            "ΑΓΙΟΙ": "ΑΓΙΩN ",
            "ΝEOI": "ΝΕΩΝ ",
            "ΝΕΕΣ": "ΝΕΩΝ ",
            "ΑΓΙΕΣ": "ΑΓΙΩΝ ",
            "ΑΡΧΑΙΑ" :"ΑΡΧΑΙΑΣ"}

   
        import itertools

        suffixes = ["AE", "IKE", "EE", "EPE", "OE","ΑΕ","ΙΚΕ","ΕΕ","ΕΠΕ","ΟΕ"]  
        letters = ['A-Z', 'Α-Ω']  # English and Greek capital letters.
        patterns = ['', '.']  # With and without dots.

            # Generate all combinations
        εταιρείες = []
        for suffix in suffixes:
            for letter_pattern in itertools.product(patterns, repeat=len(suffix)):
                new_suffix = ''.join(f"{letter}{pattern}" for letter, pattern in zip(suffix, letter_pattern))
                εταιρείες.append(new_suffix)
        
            
        sxesi_dict = {
                "ΟΦΕΙΛΕΤΗΣ"    : "ως πρωτοφειλέτη",
                "ΩΦΕΙΛΕΤΗΣ"    : "ως πρωτοφειλέτη",
                "οφειλέτης"    : "ως πρωτοφειλέτη",
                "ΠΡΩΤΟΦΕΙΛΕΤΗΣ" : "ως πρωτοφειλέτη",
                "Πρωτοφειλέτης": "ως πρωτοφειλέτη",
                "Primary Owner": "ως πρωτοφειλέτη",
                "Primary Owner": "ως πρωτοφειλέτη",
                "Εγγυητής"     : "ως εγγυητή",
                "Εγγυητής"     : "ως εγγυητή",
                "ΕΓΓΥΗΤΗΣ"     : "ως εγγυητή",
                "Guarantor"    : "ως εγγυητή",
                "Συμπιστούχος" : "ως συμπιστούχο",
                "Συνοφειλέτης" : "ως συμπιστούχο",
                "Πιστούχος"     : "ως συμπιστούχο",
                "ΠΙΣΤΟΥΧΟΣ"     : "ως συμπιστούχο",
                "Co-Owner"     : "ως συμπιστούχο",
                "Co-Owner" :   "ως συμπιστούχο"}


        def tou(string):
            if isinstance(string, str) and string.endswith(("Α","A","Η","H","Ω")):
                return "Της"
            elif isinstance(string, str) and string.endswith(("ΟΣ", "ΑΣ", "ΗΣ")):
                return "Του"
            elif isinstance(string, str) and string in female_exceptions :
                return "Της"
            elif isinstance(string, str) and string.split(' ')[-1] in εταιρείες :
                return "Της εταιρείας"
            
            else :
                return "Του"

        def ton(string):
            if isinstance(string, str) and string.endswith(("Α","A","Η","H","Ω")):
                return "Την"
            elif isinstance(string, str) and string.endswith(("ΟΣ", "ΑΣ", "ΗΣ")):
                return "Τον"
            elif isinstance(string, str) and string in female_exceptions :
                return "Την"
            elif isinstance(string, str) and string.split(' ')[-1] in εταιρείες :
                return "Την εταιρεία"
            else :
                return "Τον"

        def onoma_ait(string):
            if isinstance(string, str) and string.endswith(("ΟΣ", "ΑΣ", "ΗΣ", "A", "Η")):
                if string.endswith(("ΟΣ", "ΑΣ", "ΗΣ")) :
                    return string[:-1]  
                else :
                    return string
            else :
                return string


        def onoma_gen(string):
            if isinstance(string, str) and string.endswith("ΟΣ"):
                return string[:-1] + "Υ"
            elif isinstance(string, str) and string.endswith(("ΑΣ","ΗΣ")):
                return string[:-1]
            elif isinstance(string, str) and string.endswith(("Α","Η","Ω")):
                return string + "Σ"
            else:
                return string
        
        def eponimo_gen(string):                    
                        if isinstance(string, str) and string.endswith("ΟΣ"):
                            return string[:-1] + 'Υ'
                        elif isinstance(string, str) and string.endswith("Σ"):
                            return string[:-1] 
                        else :
                            return string

        def eponimo_ait(string):                    
            if isinstance(string, str) and string.endswith("Σ"):
                return string[:-1]
            else :
                return string

        
        def cities(string):
            city_list = []
            if 'Πόλη' in selected_columns.keys() and isinstance(string, str):
                if string in exception_dict:
                        string = exception_dict[string]
                else :
                    if isinstance(string, str) and string != "" :
                        last_two = string[-2:]
                        if last_two in ('ΑΣ', 'ΗΣ') :
                                return string[:-1]
                        elif last_two in ('ΕΣ', 'ΟΙ','ΑΙ'):
                                return string[:-2] + 'ΩΝ'
                        elif string[-1] in  ('Α','Η','A','H'):
                                return string + 'Σ'
                        elif string[-1] == 'Ι':
                                return string + 'ΟΥ'
                        elif string[-1] == 'Ο':
                                return string + 'Υ'
                        elif last_two == 'ΟΣ':
                                return string[:-1] + 'Y'
                        else :
                            return string
                            
                for key,value in replacements.items():
                    if key in string:
                        string.replace(key, value)
                return string


        associated_text_custom_ait = {
        'Άρθρο_αιτ' : "",
        'Επώνυμο_αιτ' : "",
        'Όνομα_αιτ' : "",
        Πατρώνυμο : " του ",
        Πόλη: f", {custom_city_ait}",
        "Οδός": f", {custom_street_ait} ",
        "Τ.Κ." : f", {custom_TK_ait}",
        ΑΦΜ :  f"{custom_ΑΦΜ_gen}",
        Σχέση: ",",
        # "Αριθμός" : ", με αριθμό "
            }

        def combine_columns_custom_ait(row):
            result = ""
            for k in associated_text_custom_ait.keys() :
                if k in df.columns and row[k] is not None:
                    result += associated_text_custom_ait[k] 
                    result += row[k] + " "
            return(result)
        
        associated_text_custom_gen = {
        'Άρθρο_γεν' : "",
        'Επώνυμο_γεν' : "",
        'Όνομα_γεν' : "",
        Πατρώνυμο: " του ",
        Πόλη: f", {custom_city_gen}",
        "Οδός": f",  {custom_street_gen} ",
        "Τ.Κ.":   f", {custom_TK_gen}",
        ΑΦΜ :  f"{custom_ΑΦΜ_gen}",
        Σχέση : ",",
        "Αριθμός" : ", με αριθμό "}
        # }
        def combine_columns_custom_gen(row):
            result = ""
            for k2 in associated_text_custom_gen.keys() :
                if k2 in df.columns :      
                    result += associated_text_custom_gen[k2] 
                    result += row[k2] + " "
                    
          
            return(result)
         


            
        # def SPV(val):
        #     if 'CAIRO1' in val:
        #         return df[contract]
        #     elif 'CAIRO2' in val:
        #         return df[contract]
        #     elif 'RECOVERY' in val:
        #         return df[contract]
        #     elif 'ERB' in val:
        #         return df[contract]
        #     elif 'RECOVERY' in val:
        #         return df[contract]
        #     else:
        #         return None
        

        # ----------------------------------------------Cleaning--------------------------------------------------------------------------
        def check_for_εταιρεία_gen(row):
            for col in [Όνομα, Πατρώνυμο, Επώνυμο]:
                if any(εταιρεία in row[col] for εταιρεία in εταιρείες):
                    return "Της εταιρείας"
            return None  # or some default value
        
        def check_for_εταιρεία_ait(row):
            for col in [Όνομα, Πατρώνυμο, Επώνυμο]:
                if any(εταιρεία in row[col] for εταιρεία in εταιρείες):
                    return "Της εταιρείας"
            return None  # or some default value

        def check_for_εταιρεία_edra(row):
            for col in [Όνομα, Πατρώνυμο, Επώνυμο]:
                if any(εταιρεία in row[col] for εταιρεία in εταιρείες):
                    return "η οποία εδρεύει"
            return None  # or some default value

        def categorize1(val):
            if 'SBB' in val:
                return 'επιχειρηματικού δανείου'
            elif 'MLU' in val:
                return 'στεγαστικού δανείου'
            elif 'CLB' in val:
                return 'τοκοχρεωλυτικού δανείου'
            else:
                return None
            
        def categorize2(val):
            if 'SBB' in val:
                return 'επιχειρηματικό δάνειο'
            elif 'MLU' in val:
                return 'στεγαστικό δάνειο'
            elif 'CLB' in val:
                return 'τοκοχρεωλυτικό δάνειο'
            else:
                return None

        
        if 'Επώνυμο' in selected_columns.keys():
            df["Επώνυμο_αιτ"] = df[Επώνυμο].apply(eponimo_ait) 
            df["Επώνυμο_γεν"] = df[Επώνυμο].apply(eponimo_gen)

        # # Αν υπάρχει ονοματεπώνυμο δημιουργία στήλης άρθρου/της εταιρείας
        if 'Όνομα' in selected_columns.keys() :
                df["Άρθρο_αιτ"] = df[Όνομα].apply(ton)
                df["Άρθρο_γεν"] = df[Όνομα].apply(tou)        
                df['Εταιρείας'] = df.apply(check_for_εταιρεία_gen, axis=1)
                df['Εταιρεία'] = df.apply(check_for_εταιρεία_ait, axis=1)
                df['Έδρα'] = df.apply(check_for_εταιρεία_edra, axis=1)
                df["Άρθρο_γεν"] = df.apply(lambda row: check_for_εταιρεία_gen(row) if check_for_εταιρεία_gen(row) else tou(row[Όνομα]), axis=1)
                df["Άρθρο_ αιτ"] = df.apply(lambda row: check_for_εταιρεία_ait(row) if check_for_εταιρεία_ait(row) else tou(row[Όνομα]), axis=1)        
                df["Άρθρο_γεν"] = np.where(df['Εταιρείας'] == 'Της εταιρείας', 'Της εταιρείας',df["Άρθρο_γεν"] )
                df["Άρθρο_αιτ"] = np.where(df['Εταιρεία']  == 'Την εταιρεία', 'Την εταιρεία',df["Άρθρο_αιτ"] )           
                df["ο/η"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "ο", "η")
                df["αυτόν/αυτή"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "αυτόν", "αυτή")    
                df["Όνομα_αιτ"] = df[Όνομα].apply(onoma_ait)
                df["Όνομα_αιτ"] = df["Όνομα_αιτ"].fillna('-')
                df["Όνομα_γεν"] = df[Όνομα].apply(onoma_gen)     
                df["Όνομα_αιτ"] = np.where(df['Όνομα_αιτ'] == '-',df['Επώνυμο_αιτ'],df['Όνομα_αιτ'] )
                df["Όνομα_γεν"] = np.where(df['Όνομα_γεν'] == '-',df['Επώνυμο_γεν'],df['Όνομα_γεν'] )

     
        if 'Πόλη' in selected_columns.keys() :
            df['Πόλη'] = df[Πόλη].apply(cities)
            st.write(df['Πόλη'])
    
        if 'Σχέση' in selected_columns.keys() :
            df['Σχέση'] = df['Σχέση'].replace(sxesi_dict)  

        if 'ΑΦΜ' in selected_columns.keys() :
                df[ΑΦΜ] = np.where(df[ΑΦΜ].str.len() == 10, df[ΑΦΜ].str[1:], df[ΑΦΜ])
                df[ΑΦΜ] = df[ΑΦΜ].apply(lambda x: x.zfill(9))

        if 'Πατρώνυμο' in selected_columns.keys() :
            df[Πατρώνυμο] = df[Πατρώνυμο].apply(onoma_gen)
            df[Πατρώνυμο] = df[Πατρώνυμο].fillna("-")      


        df["Στοιχεία_Αιτ_Custom"] = df.apply(combine_columns_custom_ait, axis=1).replace(" 0 ","")
        df["Στοιχεία_Γεν_Custom"] = df.apply(combine_columns_custom_gen, axis=1).replace(" 0 ","")
        df["Στοιχεία_Αιτ_Custom"] =  df["Στοιχεία_Αιτ_Custom"].apply(lambda x : str(x).replace(" , ",  ", ")) 
        df["Στοιχεία_Γεν_Custom"] =  df["Στοιχεία_Γεν_Custom"].apply(lambda x : str(x).replace(" , ",  ", ")) 


        if 'Πατρώνυμο' in selected_columns.keys() :
            df["Στοιχεία_Αιτ_Custom"] = np.where(df[Πατρώνυμο] == "-",df["Στοιχεία_Αιτ_Custom"].apply(lambda x: x.replace("   του  , κατοίκου ", ", η οποία εδρεύει στην περιοχή ")),df["Στοιχεία_Αιτ_Custom"])
            df["Στοιχεία_Γεν_Custom"] = np.where(df[Πατρώνυμο] == "-",df["Στοιχεία_Γεν_Custom"].apply(lambda x: x.replace("   του  , κατοίκου ", ", η οποία εδρεύει στην περιοχή ")),df["Στοιχεία_Γεν_Custom"])
        
        df["Στοιχεία_Αιτ_Custom"] = np.where(
        df['Εταιρείας'] == "Την εταιρεία",
        df["Στοιχεία_Αιτ_Custom"].str.replace("κατοίκου", "η οποία εδρεύει στην περιοχή").replace("   του  ",""),
        df["Στοιχεία_Αιτ_Custom"])


        df["Στοιχεία_Γεν_Custom"] = np.where(
        df['Εταιρείας'] == "Της εταιρείας",
        df["Στοιχεία_Γεν_Custom"].str.replace("κατοίκου", "η οποία εδρεύει στην περιοχή").replace("   του  ",""),
        df["Στοιχεία_Γεν_Custom"])

        # if 'Ποσό' in selected_columns.keys() :        
        #     df['Ολογράφως_Αιτ'] = df[Ποσό].apply(olog)
        #     df['Ολογράφως_Γεν'] = df[Ποσό].apply(olog_gen)

        if 'Υπόθεση' in df.columns:
            case_index = list(df.columns).index('Υπόθεση')
        else:
            case_index = 0  # Fallback index if default_column is not found

        if 'Business Unit' in df.columns:
            bu_index = list(df.columns).index('Υπόθεση')
        else:
            bu_index = 0  # Fallback index if default_column is not found


        if st.selectbox('Κατηγοριοποίηση βάσει SPV & Δανείου', options = ['Yes','No']) == 'Yes' :
            SPV = sb("SPV", options=df.columns)
            daneio= sb("Δάνειο", options=df.columns,index = bu_index)    
            contract = st.selectbox("Group column",options  = df.columns, index = case_index)
            df['placeholderdaneiou'] = df[daneio].apply(categorize1)
            df['placeholderdaneio'] =  df[daneio].apply(categorize2)
            SPV_filter = df[SPV].isin(['CAIRO1', 'CAIRO2', 'MEXICO'])
            RECOVERY_filter = df[SPV] == 'RECOVERY'
            ERB_filter = df[SPV] == 'ERB'
            df.loc[SPV_filter, 'placeholder4b'] = df[contract]
            df.loc[SPV_filter, 'placeholder4c'] = 'Σύμβαση'
            df.loc[SPV_filter | ERB_filter, 'placeholder4d'] = df['placeholderdaneiou']
            df.loc[RECOVERY_filter, 'placeholder2f'] = df[contract]
            df.loc[RECOVERY_filter, 'placeholder2g'] = 'σύμβασης'
            df.loc[df[SPV].isin(['MEXICO', 'RECOVERY']), 'placeholder2h'] = df[contract]
            df['Pool Ενεχομένων'] = df.groupby(contract)[contract].transform('count')


            sxesi_order_dict = {
                "ως πρωτοφειλέτη" : 1,
                "ως εγγυητή" : 2 ,
                "ως συμπιστούχο" : 3,

            }



            df['Order_Σχέση_Ενεχομένων'] = df['Σχέση'].replace(sxesi_order_dict)  
            
            df = df.sort_values(by=['Pool Ενεχομένων', contract,'Order_Σχέση_Ενεχομένων' ], ascending=[True, True, True])

            # \df = df.sort_values(by=[SPV, 'Pool Ενεχομένων', contract,'Σχέση' ], ascending=[True, True, True, False])
        # df = df.iloc[:,2:]
        
        extract_option = st.selectbox('Επιλογή τρόπου εξαγωγής', ['Pivot', 'Non-Pivot'])
        if extract_option ==  'Non-Pivot' :
            df.to_excel("result.xlsx", index=False)
            if st.button('Clean Data') :
                format_df('result.xlsx')
                os.startfile("result.xlsx")

        if extract_option == 'Pivot' :
          
            contract = st.selectbox("Group column",options  = df.columns, index = case_index)
            stoixeia = 'Στοιχεία_Γεν_Custom'
            df['Pool Ενεχομένων'] = df.groupby(contract)[contract].transform('count')
            # Create a unique sequence within each 'ContractID' group
            df['UniqueSeq'] = df.groupby(contract).cumcount()+1

            # Pivot the DataFrame
            df_pivot = df.pivot(index=contract, columns='UniqueSeq', values=stoixeia)

            # Reset index
            df_pivot.reset_index(inplace=True)
            
            # Rename the columns
            df_pivot.columns = [f'Stoixeia{i-1}' for i in range(1, len(df_pivot.columns)+1)]
            df_final = pd.merge(df.drop_duplicates(subset=contract), df_pivot, left_on=contract,right_on = df_pivot.iloc[:,0])
            file_name = 'pivoted.xlsx'
            df_final.to_excel(file_name,index = False)
            if st.button('Clean Data') :
                format_df(file_name)    
                os.startfile(file_name)

            if st.button('Extra Pronouns') :
                df["του/της"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "του", "της")
                df["τον/την"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "τον", "την")
                df["στον/στην"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "στον", "στην")
                df["καθ'ου/καθ'ης"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "καθ'ου", "καθ'ης")
                df["εναγόμενος/εναγόμενη"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "εναγόμενος", "εναγόμενη")
                df["εναγομένου/εναγομένης"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "εναγομένου", "εναγομένης")
                df["εναγόμενο/εναγόμενη"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "εναγόμενο", "εναγόμενη")
                df["μηνυόμενος/μηνυόμενη"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "μηνυόμενος", "μηνυόμενη")
                df["οφειλέτης/οφειλέτιδα"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "οφειλέτης", "οφειλέτιδα")
                df["οφειλέτη/οφειλέτιδας"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "οφειλέτη", "οφειλέτιδας")
                df["οφειλέτη/οφειλέτιδα"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "οφειλέτη", "οφειλέτιδα")
                df["αυτός/αυτή"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "αυτός", "αυτή")
                df["αυτού/αυτής"] = np.where(df["Άρθρο_αιτ"] == 'Τον' , "αυτού", "αυτής")
                format_df(file_name)    
                os.startfile(file_name)

if sidebar == '4) Διαχωρισμός' :
    st.header('Διαχωρισμός')
    uploaded_file = st.file_uploader("Επιλογή αρχείου Excel :")
    
    if uploaded_file is not None:
        arxiko_df=pd.read_excel(uploaded_file,sheet_name= None,dtype=str)
        selected_sheet = st.selectbox("Επιλογή φύλλου :",options = arxiko_df.keys())
        df = arxiko_df[selected_sheet] 
        options_list = ['By column',
                        'By count of column',
                        "By column 'a' and by column 'b'",
                        "By column 'a' and count of column 'b'",
                        "By column 'a' and column 'b' and count of column 'c'"
                        ,'Filter by values of a column',
                        'Pivot']
        
        split_option = st.selectbox("Select method:",options = options_list )
        
        if split_option == 'By column': 
            # 1 Replace nulls with '-'
            # 2. Group by selected column
            # 3. Create a writer object(excel file) and shorten name if it is too long
            # 4. Get each created df group 
            # 5a. Write it to a separate sheet in the writer excel file
            # 5b. Write it to a separate file and open the file explorer
            column = st.selectbox('Επιλογή Στήλης',options=df.keys())   

            df = df.fillna('-') #1
            groups = df.groupby(column) #2
            writer = pd.ExcelWriter(f"{selected_sheet}_splitted_by_{column}.xlsx", engine="xlsxwriter") #3
            option = st.selectbox('Εξαγωγή σε αρχεία ή φύλλα excel',options= ['Αρχεία','Φύλλα']) 
            if option == 'Φύλλα' :
                for name, group in groups: #4 
                    if len(name) > 25 : 
                        name = name[:25]
                    group.to_excel(writer, sheet_name=f"{name}_{len(group)}", index=False) #5a   
                    writer.save()
            # writer.close()
            if st.button('Show sheets') :
                format_df(writer)
                os.startfile(writer)
            if option == 'Αρχεία' :
                for name, group in groups: #4
                    if len(name) > 25 :  
                        name = name[:25] 
                    group.to_excel(f"{name}.xlsx",sheet_name = 'Data',index=False,engine="xlsxwriter") #5b
                    format_df(f"{name}.xlsx")
                    
                writer.save()
                # writer.close()
                if st.button('Show files') :
                    os.startfile(output)

        if split_option == "By column 'a' and by column 'b'" :      
          # 1 Replace nulls with '-'
          # 2. Group by selected column 
          # 3. Create groups list 
          # 4. Get each created group df 
          # 5. For each df in the groups list group the df into sub-dfs for the second column 
          # 6. Write each item in the final df list to a separate sheet in the same excel file
          # 7. Name it according to the first value of both columns that were used for grouping
          
            first_column = st.selectbox('Choose first column',options=df.keys())     
            second_column = st.selectbox('Choose second column',options=df.keys())     
            df[[first_column,second_column]]=df[[first_column,second_column]].fillna('-') #1 
            groups = df.groupby(first_column)  #2
            
            split_dfs = [group for _, group in groups] # 3 
            writer = pd.ExcelWriter('splitted_data.xlsx', engine='openpyxl')            
            final_dfs = []
            for d in split_dfs: #4
                sub_dfs = [g for _, g in d.groupby(second_column)] #5 
                # Create final dfs list
                final_dfs.extend(sub_dfs)  # If we used append, we would get the list of the dfs not the dfs themselves

            for f in final_dfs :
                sheet_name = f'{f[first_column].iloc[0]}_{f[second_column].iloc[0]}' #6,7
                f.to_excel(writer, sheet_name=sheet_name, index=False) 
            writer.save()
            # writer.close()
            if st.button('Split file') :
                os.startfile('splitted_data.xlsx') 

        if split_option == 'By count of column': 
            # 1. Create a new column with the count of each value in the selected ('AFM') column 
            # 2. Get a list of the unique counts in the created count column
            # 3. Create a dict compr. for each item in the list 
            #    keys   : str(df_{count})
            #    values : filter original df by count column for each item in the list 
            # 4. Write each value of the dict to a separate sheet in the same excel file (sheet_name = key of dict)
            
            column = st.selectbox('Choose a column',options=df.keys()) 
            df[column].fillna('-', inplace  = True)
            
            df['count'] = df.groupby(column)[column].transform('count') #1
            counts = df['count'].unique() #2
            result = {f'{column}_{count}': df[df['count'] == count] for count in counts} #3
        
            writer = pd.ExcelWriter(f"{selected_sheet}_splitted_by_count_of_{column}.xlsx", engine='xlsxwriter') #4
            for key, value in result.items():
                value.to_excel(writer, sheet_name=key, index=False)

            writer.save()
            # writer.close()
            if st.button('Split') :
             format_df(writer)    
             os.startfile(writer)        

        if split_option == "By column 'a' and count of column 'b'": 
            first_column = st.selectbox('Choose first column',options=df.keys())   
            second_column = st.selectbox('Choose second column',options=df.keys()) 
            df[[first_column,second_column]]=df[[first_column,second_column]].fillna('-')
            writer = pd.ExcelWriter('splitted_by_count_dfs.xlsx', engine='xlsxwriter')
            
            splitted_dfs = [g for name, g in df.groupby(first_column)]
            for i,s in enumerate (splitted_dfs) :
                s['Count'] = s.groupby(second_column)[second_column].transform('count')
                counts = s['Count'].unique()
                result = {f"{s[first_column].iloc[0]}_{count}_{len(s)}": s[s['Count'] == count] for count in counts}
                for key, value in result.items():     
                    value.to_excel(writer, sheet_name=key, index=False)
            writer.save() 
            # writer.close()
            if st.button('Split') :
                os.startfile('splitted_by_count_dfs.xlsx')

        if split_option == "By column 'a' and column 'b' and count of column 'c'": 
            first_column = st.selectbox('Choose first column to filter by',options=df.keys())   
            second_column = st.selectbox('Choose second column to filter by',options=df.keys()) 
            third_column = st.selectbox('Choose third column (count)',options=df.keys()) 

            df[[first_column,second_column,third_column]]=df[[first_column,second_column,third_column]].fillna('-')
            groups = df.groupby(first_column)
            split_dfs = [g for _, g in groups] 
            final_dfs = []
            for d in split_dfs:
                sub_dfs = [g for _, g in d.groupby(second_column)]
                final_dfs.extend(sub_dfs)          
            writer = pd.ExcelWriter('splitted_3.xlsx', engine='openpyxl')
            st.write(len(final_dfs))

            for i,f in enumerate(final_dfs) :
                f['Count'] = f.groupby(third_column)[third_column].transform('count')
                counts = sorted(f['Count'].unique())
                result = {f'{f[first_column].iloc[0]}_{count}_{f[second_column].iloc[0]}': f[f['Count'] == count] for count in counts}
                option = st.selectbox('Pivoted or no', options = ['Pivoted', 'Unpivoted'])
                for k, v in result.items():
                    
                    if option == 'Pivoted' :
                        pivoted_df = pd.pivot_table(v, index=third_column, columns=v.groupby(third_column).cumcount()+1, values= v.columns ,aggfunc='first')
                        pivoted_df.columns = [f'{col}' for col in pivoted_df.columns] # Redunadnt
                        pivoted_df = pivoted_df.reset_index()
                        pivoted_df.to_excel(writer, sheet_name=k, index=False)
                        
                    if option == 'Unpivoted' :
                        for k, v in result.items():
                            v.to_excel(writer, sheet_name=k, index=False)

                if st.button('Split') :
                    writer.save()
                    # writer.close() 
                    format_df(writer)    
                    os.startfile(writer)
                    
        if split_option == 'Filter by values of a column' : 
            # 1. Create a slider to define number of values to filter by 
            # 2. Choose column to filter by
            # 3. Select values 
            # 4. Append them to a list
            # 4. Use 'isin' (list -> column) 

            slider = st.slider('Select number of values to filter by',1,5) #1
            column = st.selectbox('Choose a column',options=df.keys()) #2
            filter_values = []
            for i in range(slider):
                filter_values.append(st.selectbox(f'Choose value {i+1}', options = set(df[column]))) #3,4
            final_df = df[df[column].isin(filter_values)] #5
            if st.button('Filter') :
                file_name = 'filtered.xlsx'
                final_df.to_excel(file_name)
                format_df(file_name)    
                os.startfile(file_name)
                    
        if split_option == 'Pivot' :
            contract = 'CASE CONTR NUM'
            stoixeia = 'Στοιχεία_Γεν_Custom'
            df['Pool Ενεχομένων'] = df.groupby(contract)[contract].transform('count')
            # test = df[['Pool Ενεχομένων','Στοιχεία_Γεν_Custom','CASE CONTR NUM']]
            # Create a unique sequence within each 'ContractID' group
            df['UniqueSeq'] = df.groupby(contract).cumcount()+1

            # Pivot the DataFrame
            df_pivot = df.pivot(index=contract, columns='UniqueSeq', values=stoixeia)

            # Reset index
            df_pivot.reset_index(inplace=True)

            # Rename the columns
            df_pivot.columns = [contract] + [f'Stoixeia{i}' for i in range(1, len(df_pivot.columns))]
            file_name = 'pivoted.xlsx'
            df_pivot.to_excel(file_name,index = False)
            format_df(file_name)    
            os.startfile(file_name)

if sidebar == '2) Συμπλήρωση Excel βάσει δύο στηλών' :  #✅
    # Πώς : np.where(main+extra_cols_dict[v] (= If extra_cols are null),
            # main+extra_cols_dict[k]( = then insert main_cols)
            # main+extra_cols_dict[v]( = else insert extra_cols)
    st.header("Συμπλήρωση Excel")
    main_list =['main_name','main_surname','main_father_name',
    'main_afm',
    'main_street', 
    'main_tk',
    'main_city',
    'main_rt'
    ]

    extra_list = [
    'extra_name', 
    'extra_surname', 
    'extra_father_name', 
    'extra_afm',
    'extra_street', 
    'extra_tk',
    'extra_city', 
    'extra_rt'] 

    uploaded_file = st.file_uploader("Επιλογή αρχείου")
    if uploaded_file is not None:
        arxiko_df=pd.read_excel(uploaded_file,sheet_name= None,dtype=str)
        selected_sheet = st.selectbox("Επιλογή φύλλου :",options = arxiko_df.keys())
        df = arxiko_df[selected_sheet]
        
        options = st.selectbox("Επιλογές", ['Αυτόματη Εύρεση Στηλών', 'Χειροκίνητη Εύρεση Στηλών'])
        col1, col2= st.columns(2)

        if options == 'Χειροκίνητη Εύρεση Στηλών' :
            with col1:
                st.header("Κύρια Στοιχεία")
                m_columns = []        
                for m in main_list :
                    m = st.selectbox(m,options=df.columns)
                    m_columns.append(m)

            with col2 :
                st.header('Έξτρα Στοιχεία')
                e_columns= []
                for e in extra_list :
                    e = st.selectbox(e,options=df.columns)
                    e_columns.append(e)
            
            if st.button('Select columns'):
                total_dict = dict(zip(m_columns, e_columns))
                if "Unnamed: 0" in total_dict:
                    del total_dict["Unnamed: 0"]
                for k, v in total_dict.items():
                    df[f"total_{k}"] = np.where(df[v].isna(), df[k], df[v])
                df.iloc[:,-8:]
                df.to_excel('total.xlsx')
                os.startfile('total.xlsx')

        if options == 'Αυτόματη Εύρεση Στηλών' :
            if st.button('Select columns'):
                            m_auto_columns  = ['main CUST FIRST NAME',
                                          'main CUST LAST NAME',
                                          'main CUST FATH NAME',
                                          'main CUST AFM',
                                          'STREET',
                                          'POSTCODE',
                                          'CITY']
                            e_auto_columns  = ['EXTRA FIRST NAME',
                                               'EXTRA LAST NAME',
                	                           'EXTRA FATH NAME',
                                               'EXTRA AFM',
                                               'EXTRA STREET',
                                               'EXTRA POSTCODE',
                                               'EXTRA CITY'                                        
                                          ]
                            total_dict = dict(zip(m_auto_columns, e_auto_columns))
                            if "Unnamed: 0" in total_dict:
                                del total_dict["Unnamed: 0"]
                            for k, v in total_dict.items():
                                df[f"total_{k}"] = np.where(df[v].isna(), df[k], df[v])
                            df.iloc[:,-8:]
                            df.to_excel('total.xlsx')
                            os.startfile('total.xlsx')





if sidebar == '1) Συνένωση Excel' :
 col1,col2,col3 = st.columns(3)
 st.subheader('Συνένωση Αρχείων Excel:')
 uploaded_file_1 = st.file_uploader("Παρακαλώ επιλέξτε το πρώτο αρχείο :")
 uploaded_file_2 = st.file_uploader("Παρακαλώ επιλέξτε το δεύτερο αρχείο :")

 if uploaded_file_1 is not None and uploaded_file_2 is not None:
    arxiko_df_1 = pd.read_excel(uploaded_file_1,sheet_name= None,dtype=str)
    arxiko_df_2 = pd.read_excel(uploaded_file_2,sheet_name= None,dtype=str)
    selected_sheet_1 = st.selectbox("Παρακαλώ επιλέξτε το πρώτο φύλλο :",options = arxiko_df_1.keys())
    selected_sheet_2 = st.selectbox("Παρακαλώ επιλέξτε το δεύτερο φύλλο :",options = arxiko_df_2.keys())    
    df_1 = arxiko_df_1[selected_sheet_1]
    df_2 = arxiko_df_2[selected_sheet_2]
    if st.button('Προβολή επιλεγμένου φύλλου πρώτου αρχείου') : 
        df_1
    if st.button('Προβολή επιλεγμένου φύλλου δεύτερου αρχείου') : 
        df_2
    key_1 = st.selectbox('Παρακαλώ επιλέξτε την πρώτη στήλη-κλειδί', options = df_1.columns)
    key_2 = st.selectbox('Παρακαλώ επιλέξτε την δεύτερη στήλη-κλειδί', options = df_2.columns) 
    merge_list = ['Οριζόντια Συνένωση (Left Join)',
                  'Οριζόντια Συνένωση μόνο των κοινών εγγραφών (Inner Join)',
                  'Κάθετη Συνένωση',
                  'Οριζόντια Συνένωση χωρίς κλειδί (Concatenate)',
                  'Vlookup(first match)']
    merge_button = sb('Επιλογή μεθόδου συνένωσης', options = merge_list)


    if merge_button == 'Οριζόντια Συνένωση (Left Join)' :
        merged = pd.merge(df_1,df_2,left_on = key_1,right_on=key_2,how='left')
        st.dataframe(merged)
        merged.to_excel('left_merged.xlsx',index=False)
        os.startfile('left_merged.xlsx')
    
    if merge_button == 'Οριζόντια Συνένωση μόνο των κοινών εγγραφών (Inner Join)' :
        merged = pd.merge(df_1,df_2,left_on = key_1,right_on=key_2,how='inner')
        st.dataframe(merged)
        merged.to_excel('inner_merged.xlsx',index=False)
        os.startfile('inner_merged.xlsx')
    
    if merge_button == 'Κάθετη Συνένωση' :
        synenomeno = pd.concat([df_1,df_2],axis=0,ignore_index = True)
        synenomeno.to_excel('vertical_merged.xlsx',index=False)
        os.startfile('vertical_merged.xlsx')
    
    if merge_button == 'Οριζόντια Συνένωση χωρίς κλειδί (Concatenate)' :
        synenomeno = pd.concat([df_1,df_2],axis=1,ignore_index = True)
        synenomeno.to_excel('horizontal_merged.xlsx',index=False)
        os.startfile('horizontal_merged.xlsx')
    
    if merge_button == 'Vlookup(first match)' :
        synenomeno = pd.merge(df_1,df_2.drop_duplicates(key_2),left_on = key_1,right_on=key_2,how='left')
        synenomeno.to_excel('vlookup.xlsx', index=False)
        os.startfile('vlookup.xlsx')
    
    # if merge_button == 'Συρραφή όλων των αρχείων στον φάκελο' :
    #     pass        
    #     files = glob.glob(os.path.join(fd.askdirectory(), "xl*"))s
    #     for f in files :
    #         df = pd.read_excel(f,dtype=str)
    #         df_list.append(df)

    #     synenomeno = pd.concat([df_list],axis=1,ignore_index = True)
    #     synenomeno.to_excel('horizontal_merged.xlsx')
    #     os.startfile('horizontal_merged.xlsx')


if sidebar == '6) Μαζικές εργασίες εγγράφων' :

    # 1. Configure logging
    # 2. Select tempate doc, excel data to be spillted and txt file with placeholders
    # 3. Rename the columns
    # 4. Group the df to separate sheets in a new excel file
    # 5. Format Word
    # 6. Replace placeholders

    from pathlib import Path
    import tempfile
    
    RESULTS =  Path(r"C:\Users\pallist\Desktop\ΤΡΕΧΟΝΤΑ\Testing Folder\Results")

    # 1. Configure logging
    x = os.chdir(RESULTS)
    LOG_FILENAME = os.path.join(RESULTS,'Logfile.log')
    LOG_LEVEL = logging.DEBUG
    logging.basicConfig(filename=LOG_FILENAME, level=LOG_LEVEL, format='%(asctime)s - %(levelname)s - %(message)s')
    logging.info("Starting script")

    # 2.Define a function to create a temporary copy of the selected file to  get its path to use it later for the template,excel and placeholders
    def file_upload_prompt(prompt):
        uploaded_file = st.file_uploader(prompt)
        if uploaded_file is None:
            logging.error(f'No file selected for {prompt}.')
            return None
        file_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
        with open(file_path, 'wb') as f:
            f.write(uploaded_file.getbuffer())   
        logging.info(f'{uploaded_file.name} uploaded and saved to {file_path}.')
        return file_path


    excel_file_path = file_upload_prompt('Select the data to be splitted')
    template_file_path = file_upload_prompt('Select the template document')
    placeholders_path = file_upload_prompt('Select the text file containing the list of placeholders')
    
    st.subheader('File list :') 
    st.write(excel_file_path)
    st.write(template_file_path)
    st.write(placeholders_path)



    # 3. Rename columns to match with template and list of placeholders 
    COLUMNS_TO_RENAME = {
        'Σύμβαση': 'placeholdersym',
        'Ολογράφως_Αιτ': 'placeholderolog',
        'Stoixeia1': 'placeholderstoixeia1',    
        'Stoixeia2' :    'placeholderstoixeia2',
        'Stoixeia3' :    'placeholderstoixeia3',
        'Stoixeia4' :'placeholderstoixeia4',
        'Denounced Date': 'placeholderdate',
        'Denounced Amount' : 'placeholderamount',
        'Case Currency' : 'placeholdercur'  
    }


    # 4. Group the df to separate sheets in a new excel file
    st.write(excel_file_path)
    df = pd.read_excel(excel_file_path, engine='openpyxl')


    auto_options = ['Frontier','Non Frontier']

    category_option = st.selectbox('Επιλογή Κατηγορίας Εξωδίκων', options = auto_options)
    file_name_option = st.selectbox('Επιλογή Ονομασίας Αρχείων', options = df.columns)

    if category_option == 'Non Frontier' :
    
        def split_dataframe_to_excel_sheets_and_copy_templates(df):
            to_count = st.selectbox('Select a column to split by the count of it (Pool Ενεχομένων)',list(df.columns))
            group = st.selectbox('Select a column to split by the values of it (SPV)',list(df.columns))
            df_grouped = df.groupby([group, to_count])
        
            
            with pd.ExcelWriter(excel_file_path, engine='openpyxl', mode='a') as writer:
                for name, group in df_grouped:
                    group_name = '_'.join(str(n) for n in name)
                    group.to_excel(writer, sheet_name=group_name, index=False)
                    destination_file = os.path.join(RESULTS, f"{group_name}.docx")  
                    shutil.copy(template_file_path, destination_file)
        split_dataframe_to_excel_sheets_and_copy_templates(df)
        spv_column = 'ΧΑΡΤΟΦΥΛΑΚΙΟ'
        
        def save_new_doc(new_doc, sheet_name, row, index, spv_column):
            """Save a new Document object to a file."""
            folder_path = os.path.join(os.getcwd(), sheet_name)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
            new_doc_path = os.path.join(folder_path, f"{sheet_name}_{row[file_name_option]}_{index+1}.docx")
            
            new_doc.save(new_doc_path)
            print(f"Saved document to {new_doc_path}")

        def generate_documents(template_files, excel_file, placeholders): 
            """Generate the documents based on templates and data from an Excel file."""
            
            for template_path in template_files:
                try:
                    sheet_name = os.path.basename(template_path).split(".")[0]
                    df = load_dataframe(excel_file, sheet_name)
            
                    total_rows = len(df)
                    df.apply(lambda row: replace_placeholders_and_save_doc(row.name, row, template_path, sheet_name, total_rows, placeholders,'ΧΑΡΤΟΦΥΛΑΚΙΟ'), axis=1) 

                except Exception as e:
                    logging.error(f"Failed to process template {template_path} with error {str(e)}")
    
        def replace_placeholders_and_save_doc(index, row, template_path, sheet_name, total_rows, placeholders, spv_column_var):
        
        
            # st.write(f"Processing row {index+1} of {total_rows} in sheet {sheet_name}")
            try:
                logging.debug('Starting replace_placeholders_and_save_doc function.')
                doc = Document(template_path) 
                new_doc = Document()
                for paragraph in doc.paragraphs:
                    new_paragraph = new_doc.add_paragraph()
                    copy_paragraph_formatting(paragraph, new_paragraph)
                    for run in paragraph.runs:
                        replace_placeholders_in_run(run, row, placeholders, new_paragraph)
                        
                logging.info(f"Processing {sheet_name}: {index + 1}/{total_rows}")
                save_new_doc(new_doc, sheet_name, row, index,'ΧΑΡΤΟΦΥΛΑΚΙΟ')
                # Non Frontier -> save_new_doc(new_doc, sheet_name, row, index, spv_column)
                
                
            except Exception as e:
                logging.error(f"Failed to process {sheet_name}: {index + 1}/{total_rows} with error {str(e)}")
                st.write(f"Failed to process {sheet_name}: {index + 1}/{total_rows} with error {str(e)}")

        
    elif category_option == 'Frontier' :

        def split_dataframe_to_excel_sheets_and_copy_templates_simple(df):
            to_count = st.selectbox('Select a column to split by the count of it (Pool Ενεχομένων)',list(df.columns))
            df_grouped = df.groupby(to_count)
        
            
            with pd.ExcelWriter(excel_file_path, engine='openpyxl', mode='a') as writer:
                for name, group in df_grouped:
                
                    group_name = str(name) + "_ΕΝΕΧΟΜΕΝΟΙ" 
                    st.write(group_name)
                    group.to_excel(writer, sheet_name=group_name, index=False)
                    destination_file = os.path.join(RESULTS, f"{group_name}.docx")  
                    shutil.copy(template_file_path, destination_file)

        split_dataframe_to_excel_sheets_and_copy_templates_simple(df)
  
        
        def save_new_doc(new_doc, sheet_name, row, index):
            """Save a new Document object to a file."""
            folder_path = os.path.join(os.getcwd(), sheet_name)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)
            new_doc_path = os.path.join(folder_path, f"{sheet_name}_{row[file_name_option]}_{index+1}.docx")
            
            new_doc.save(new_doc_path)
            print(f"Saved document to {new_doc_path}")


        def generate_documents(template_files, excel_file, placeholders): 
            """Generate the documents based on templates and data from an Excel file."""
            
            for template_path in template_files:
                try:
                    sheet_name = os.path.basename(template_path).split(".")[0]
                    df = load_dataframe(excel_file, sheet_name)
            
                    total_rows = len(df)
                    df.apply(lambda row: replace_placeholders_and_save_doc(row.name, row, template_path, sheet_name, total_rows, placeholders), axis=1) 

                except Exception as e:
                    logging.error(f"Failed to process template {template_path} with error {str(e)}")


        def replace_placeholders_and_save_doc(index, row, template_path, sheet_name, total_rows, placeholders):
            
            
                # st.write(f"Processing row {index+1} of {total_rows} in sheet {sheet_name}")
                try:
                    logging.debug('Starting replace_placeholders_and_save_doc function.')
                    doc = Document(template_path) 
                    new_doc = Document()
                    for paragraph in doc.paragraphs:
                        new_paragraph = new_doc.add_paragraph()
                        copy_paragraph_formatting(paragraph, new_paragraph)
                        for run in paragraph.runs:
                            replace_placeholders_in_run(run, row, placeholders, new_paragraph)
                            
                    logging.info(f"Processing {sheet_name}: {index + 1}/{total_rows}")
                    save_new_doc(new_doc, sheet_name, row, index)
                    # Non Frontier -> save_new_doc(new_doc, sheet_name, row, index, spv_column)
                    
                    
                except Exception as e:
                    logging.error(f"Failed to process {sheet_name}: {index + 1}/{total_rows} with error {str(e)}")
                    st.write(f"Failed to process {sheet_name}: {index + 1}/{total_rows} with error {str(e)}")

            

    #5 . Format Word

    def copy_paragraph_formatting(src_paragraph, dest_paragraph):
        """Copy the formatting from src_paragraph to dest_paragraph.""" 
        attrs = ['alignment', 'left_indent', 'right_indent', 'first_line_indent', 
                'space_before', 'space_after', 'line_spacing']
        for attr in attrs:
            setattr(dest_paragraph.paragraph_format, attr, getattr(src_paragraph.paragraph_format, attr))

    def copy_run_attributes(src_run, dest_run):
        """Copy font attributes from src_run to dest_run.""" 
        attrs = ['bold', 'italic', 'underline', 'strike', 'double_strike', 
                'all_caps', 'size', 'color.rgb', 'highlight_color', 'name']
        for attr in attrs:
            if '.' in attr:
                attr, sub_attr = attr.split('.')
                setattr(getattr(dest_run.font, attr), sub_attr, getattr(getattr(src_run.font, attr), sub_attr))
            else:
                setattr(dest_run.font, attr, getattr(src_run.font, attr))

    # 6. Replace placeholders 

    def get_placeholders_from_file(path):
        """Read placeholders from a text file.""" 
        with open(path) as f:
            return f.read().splitlines()

    def load_excel_file(path): 
        """Load an Excel file into a pandas ExcelFile object."""
        return pd.ExcelFile(path)

    def get_template_files(path): 
        """Get a list of all .docx files in a directory.""" 
        return glob.glob(os.path.join(path, "*.docx"))

    def load_dataframe(excel_file, sheet_name_var):
        """Load a DataFrame from an Excel file."""
        df = pd.read_excel(excel_file,sheet_name= sheet_name_var, dtype=str)
        df = df.fillna("")
        df.rename(columns=COLUMNS_TO_RENAME, inplace=True)
        
        return df


    def replace_placeholders_in_run(run, row, placeholders, new_paragraph):
        """Replace placeholders in a run of text."""
        for placeholder in placeholders:
            if placeholder not in row:
                print(f"Warning: Placeholder {placeholder} not found in row")
            pattern = re.escape(placeholder)
            if re.search(pattern, run.text):
                run.text = re.sub(pattern, str(row[placeholder]), run.text) if placeholder in row else run.text
        run.text = re.sub(r'(\b[A-Za-z]\b)\s*&\s*(\b[A-Za-z]\b)', r'\1 \2', run.text)
        new_run = new_paragraph.add_run(run.text)
        copy_run_attributes(run, new_run)


    placeholders = get_placeholders_from_file(placeholders_path)
    excel_file = load_excel_file(excel_file_path)
    template_files = get_template_files(RESULTS)

    if st.button('Generate Documents') :
        generate_documents(template_files, excel_file, placeholders)
        st.success(f"Generated Documents inm {os.getcwd()}")

if sidebar =='Αρχική Σελίδα':
    st.title("Theo's Formatting Tool (TFT)  🎉")
